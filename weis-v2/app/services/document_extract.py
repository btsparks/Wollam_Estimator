"""Document text extraction for uploaded PM context documents.

Handles PDF, Excel (.xlsx/.xls), CSV, and plain text files.
Returns extracted text suitable for AI analysis.
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Max characters to extract per document (prevents token overflow)
MAX_TEXT_LENGTH = 80_000


def extract_text(filepath: Path) -> str:
    """Extract text content from a document file.

    Supports: .pdf, .xlsx, .xls, .csv, .txt
    Returns extracted text (truncated to MAX_TEXT_LENGTH).
    """
    suffix = filepath.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(filepath)
    elif suffix in (".xlsx", ".xls"):
        return _extract_excel(filepath)
    elif suffix == ".csv":
        return _extract_csv(filepath)
    elif suffix in (".txt", ".text", ".md"):
        return _extract_text(filepath)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(filepath: Path) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    pages = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages.append(f"--- Page {i + 1} ---\n{text}")

            # Also try extracting tables
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                if table:
                    table_text = _format_table(table)
                    if table_text:
                        pages.append(f"[Table {t_idx + 1} on page {i + 1}]\n{table_text}")

    result = "\n\n".join(pages)
    return result[:MAX_TEXT_LENGTH]


def _extract_excel(filepath: Path) -> str:
    """Extract text from Excel workbook."""
    import openpyxl

    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheets = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Convert each cell to string, skip fully empty rows
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append("\t".join(cells))

        if rows:
            sheets.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))

    wb.close()
    result = "\n\n".join(sheets)
    return result[:MAX_TEXT_LENGTH]


def _extract_csv(filepath: Path) -> str:
    """Extract text from CSV file."""
    rows = []
    with open(filepath, "r", encoding="utf-8-sig", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            if any(c.strip() for c in row):
                rows.append("\t".join(row))

    result = "\n".join(rows)
    return result[:MAX_TEXT_LENGTH]


def _extract_text(filepath: Path) -> str:
    """Extract text from plain text file."""
    text = filepath.read_text(encoding="utf-8", errors="replace")
    return text[:MAX_TEXT_LENGTH]


def _extract_docx(filepath: Path) -> str:
    """Extract text from Word .docx files using python-docx.

    Note: .doc (old binary format) is not supported by python-docx.
    We attempt it anyway — it will fail gracefully for true .doc files.
    """
    import docx

    try:
        doc = docx.Document(filepath)
    except Exception:
        raise ValueError(f"Could not read Word file (may be old .doc format): {filepath.name}")

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append("\t".join(cells))

    result = "\n".join(paragraphs)
    return result[:MAX_TEXT_LENGTH]


def _format_table(table: list[list]) -> str:
    """Format a pdfplumber table as tab-separated text."""
    rows = []
    for row in table:
        cells = [str(c) if c is not None else "" for c in row]
        if any(c.strip() for c in cells):
            rows.append("\t".join(cells))
    return "\n".join(rows)
