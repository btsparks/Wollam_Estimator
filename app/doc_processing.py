"""WEIS document processing for active bid documents.

Extracts text from PDF, DOCX, XLSX, MD, and TXT files,
then chunks the text into searchable segments (~1500 words each).
"""

import re


def extract_document(file_bytes: bytes, filename: str) -> dict:
    """Extract text from a document file.

    Routes to the correct extractor based on file extension.

    Args:
        file_bytes: Raw file content.
        filename: Original filename (used to determine file type).

    Returns:
        dict with keys:
            text: Extracted text content.
            page_count: Number of pages (if applicable).
            word_count: Total word count.
            status: 'success', 'partial', or 'failed'.
            warning: Warning message (if any).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "xlsx": _extract_xlsx,
        "md": _extract_text,
        "txt": _extract_text,
    }

    extractor = extractors.get(ext)
    if not extractor:
        return {
            "text": "",
            "page_count": None,
            "word_count": 0,
            "status": "failed",
            "warning": f"Unsupported file type: .{ext}",
        }

    try:
        return extractor(file_bytes, filename)
    except Exception as e:
        return {
            "text": "",
            "page_count": None,
            "word_count": 0,
            "status": "failed",
            "warning": f"Extraction error: {e}",
        }


def _extract_pdf(file_bytes: bytes, filename: str) -> dict:
    """Extract text from PDF using PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    empty_pages = 0

    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text}")
        else:
            empty_pages += 1

    doc.close()

    page_count = len(pages) + empty_pages
    full_text = "\n\n".join(pages)
    word_count = len(full_text.split())

    warning = None
    status = "success"

    if empty_pages > 0 and not pages:
        status = "failed"
        warning = (
            f"All {page_count} pages appear to be scanned images with no extractable text. "
            "OCR is not currently supported."
        )
    elif empty_pages > 0:
        status = "partial"
        warning = (
            f"{empty_pages} of {page_count} pages had no extractable text "
            "(likely scanned images)."
        )

    return {
        "text": full_text,
        "page_count": page_count,
        "word_count": word_count,
        "status": status,
        "warning": warning,
    }


def _extract_docx(file_bytes: bytes, filename: str) -> dict:
    """Extract text from Word document using python-docx."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n\n".join(parts)
    word_count = len(full_text.split())

    return {
        "text": full_text,
        "page_count": None,
        "word_count": word_count,
        "status": "success" if full_text.strip() else "failed",
        "warning": None if full_text.strip() else "No text content found in document.",
    }


def _extract_xlsx(file_bytes: bytes, filename: str) -> dict:
    """Extract text from Excel spreadsheet using openpyxl."""
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_lines = [f"=== Sheet: {sheet_name} ==="]

        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip()
            if line.replace("|", "").strip():
                sheet_lines.append(line)

        if len(sheet_lines) > 1:  # more than just the header
            parts.append("\n".join(sheet_lines))

    wb.close()

    full_text = "\n\n".join(parts)
    word_count = len(full_text.split())

    return {
        "text": full_text,
        "page_count": len(wb.sheetnames) if parts else None,
        "word_count": word_count,
        "status": "success" if full_text.strip() else "failed",
        "warning": None if full_text.strip() else "No data found in spreadsheet.",
    }


def _extract_text(file_bytes: bytes, filename: str) -> dict:
    """Extract text from MD or TXT files (direct UTF-8 decode)."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    word_count = len(text.split())

    return {
        "text": text,
        "page_count": None,
        "word_count": word_count,
        "status": "success" if text.strip() else "failed",
        "warning": None if text.strip() else "File is empty.",
    }


# ---------------------------------------------------------------------------
# Text Chunking
# ---------------------------------------------------------------------------

# Patterns that suggest section boundaries
_SECTION_PATTERNS = re.compile(
    r"^(?:"
    r"---\s*Page\s+\d+\s*---"        # PDF page markers
    r"|===\s*Sheet:.*==="              # Excel sheet markers
    r"|#{1,4}\s+.+"                    # Markdown headings
    r"|SECTION\s+\d+"                  # Spec section markers
    r"|ARTICLE\s+\d+"                  # Spec article markers
    r"|PART\s+\d+"                     # Spec part markers
    r"|DIVISION\s+\d+"                 # CSI division markers
    r")",
    re.MULTILINE | re.IGNORECASE,
)


def chunk_text(text: str, max_words: int = 1500) -> list[dict]:
    """Split text into searchable chunks of approximately max_words.

    Splits on section headings, page markers, or paragraph boundaries.

    Args:
        text: Full extracted text.
        max_words: Target maximum words per chunk.

    Returns:
        List of dicts with keys: chunk_index, chunk_text, section_heading.
    """
    if not text or not text.strip():
        return []

    # Split text into sections at recognized boundaries
    sections = _split_into_sections(text)

    chunks = []
    chunk_index = 0

    for heading, section_text in sections:
        words = section_text.split()

        if len(words) <= max_words:
            # Section fits in one chunk
            if section_text.strip():
                chunks.append({
                    "chunk_index": chunk_index,
                    "chunk_text": section_text.strip(),
                    "section_heading": heading,
                })
                chunk_index += 1
        else:
            # Split large section into sub-chunks on paragraph boundaries
            paragraphs = re.split(r"\n\s*\n", section_text)
            current_chunk = []
            current_words = 0

            for para in paragraphs:
                para_words = len(para.split())

                if current_words + para_words > max_words and current_chunk:
                    chunk_text_str = "\n\n".join(current_chunk).strip()
                    if chunk_text_str:
                        chunks.append({
                            "chunk_index": chunk_index,
                            "chunk_text": chunk_text_str,
                            "section_heading": heading,
                        })
                        chunk_index += 1
                    current_chunk = []
                    current_words = 0

                current_chunk.append(para)
                current_words += para_words

            # Remaining text
            if current_chunk:
                chunk_text_str = "\n\n".join(current_chunk).strip()
                if chunk_text_str:
                    chunks.append({
                        "chunk_index": chunk_index,
                        "chunk_text": chunk_text_str,
                        "section_heading": heading,
                    })
                    chunk_index += 1

    return chunks


def _split_into_sections(text: str) -> list[tuple[str | None, str]]:
    """Split text into (heading, content) pairs at section boundaries."""
    matches = list(_SECTION_PATTERNS.finditer(text))

    if not matches:
        return [(None, text)]

    sections = []

    # Text before first heading
    if matches[0].start() > 0:
        before = text[: matches[0].start()]
        if before.strip():
            sections.append((None, before))

    for i, match in enumerate(matches):
        heading = match.group(0).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end]
        if content.strip():
            sections.append((heading, content))

    return sections
